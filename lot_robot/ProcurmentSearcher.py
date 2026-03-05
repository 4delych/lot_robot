import os
import sys
import time
import re
import logging
import zipfile
import html
import tempfile
import subprocess
import shutil
import locale
from urllib.parse import urljoin, urlparse, parse_qs, unquote
import re
from urllib.parse import urlparse
import json

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from bs4 import BeautifulSoup
from io import BytesIO

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    import xlrd
except ImportError:
    xlrd = None

try:
    import pythoncom  # type: ignore
except ImportError:
    pythoncom = None  # type: ignore

try:
    import win32com.client as win32_client  # type: ignore
    from win32com.client import constants as win32_constants  # type: ignore
except ImportError:
    win32_client = None  # type: ignore
    win32_constants = None

from config import CONFIG, PURCHASE_STAGES, LAWS
from procurement_sources import ProcurementSource, ZakupkiGovSource, TektorgSource, BidzaarSource

logger = logging.getLogger(__name__)


class ProcurementSearcher:
    """Handle web scraping operations with proper error handling and session management."""

    _ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}
    _TEXT_EXTRACT_EXTENSIONS = {
        ".pdf",
        ".doc",
        ".docx",
        ".xls",
        ".xlsx",
        ".xlsm",
        ".xltx",
        ".xltm",
        ".txt",
        ".rtf",
        ".odt",
        ".ods",
    }
    _ARCHIVE_MAX_DEPTH = 2
    _ARCHIVE_MAX_FILES = 100
    _ARCHIVE_MAX_TOTAL_SIZE = 50 * 1024 * 1024
    _TZ_STRONG_MARKERS = (
        "техническое задание",
        "тех задание",
        "техзадание",
        "technical specification",
        "statement of work",
    )
    _TZ_WEAK_MARKERS = (
        " тз ",
        "/тз/",
        "\\тз\\",
        "(тз)",
        "_тз",
        "тз_",
        "-тз",
        "тз-",
    )

    def build_lot_documents_text(self, documents: list[dict], max_chars: int = 120_000) -> str:
        """
        Склеивает текст из всех документов в один текст.
        Использует существующий парсер _extract_text_from_content().
        Ограничивает общий размер, чтобы не убиться об лимиты LLM.
        """
        parts: list[str] = []
        total = 0
        ordered_documents = sorted(documents, key=self._document_priority_sort_key)

        logger.info("Порядок обхода документов перед объединением:")
        for idx, doc in enumerate(ordered_documents, start=1):
            path_hint = self._document_path_hint(doc)
            logger.info(
                "  %s) приоритет=%s файл=%r",
                idx,
                self._document_priority_score(doc),
                path_hint,
            )

        for doc in ordered_documents:
            filename = self._determine_document_filename(doc)
            raw_text = self._extract_text_from_content(
                doc.get("content") or b"",
                filename,
                doc.get("content_type", "") or "",
            )

            print(f"[DOC TEXT] {doc.get('name') or filename} -> {len(raw_text or '')} chars")

            # ✅ сохраняем переносы строк, чистим пробелы построчно
            raw = (raw_text or "").replace("\r\n", "\n").replace("\r", "\n")
            lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in raw.split("\n")]
            lines = [ln for ln in lines if ln]  # убираем пустые строки
            text = "\n".join(lines).strip()

            if not text or self._looks_like_garbage_text(text):
                logger.info(
                    "Документ пропущен: пустой или мусорный текст: файл=%r",
                    doc.get("name") or filename,
                )
                continue

            if doc.get("source_archive"):
                logger.info(
                    "Документ из архива добавлен в объединенный текст: архив=%r файл=%r символов=%s",
                    doc.get("source_archive"),
                    doc.get("name") or filename,
                    len(text),
                )
            else:
                logger.info(
                    "Документ добавлен в объединенный текст: файл=%r символов=%s",
                    doc.get("name") or filename,
                    len(text),
                )

            header = f"\n\n===== ДОКУМЕНТ: {doc.get('name') or filename} =====\n"
            chunk = header + text

            if total + len(chunk) > max_chars:
                logger.info(
                    "Документ не полностью вошел в объединенный текст из-за лимита: файл=%r доступно_символов=%s размер_фрагмента=%s",
                    doc.get("name") or filename,
                    max_chars - total,
                    len(chunk),
                )
                remain = max_chars - total
                if remain > 0:
                    parts.append(chunk[:remain])
                parts.append("\n\n[ТЕКСТ ОБРЕЗАН ПО ЛИМИТУ]")
                break

            parts.append(chunk)
            total += len(chunk)

        combined = "".join(parts).strip()

        # ✅ DEBUG: проверяем, что реально склеилось
        preview_len = 800
        print("\n===== COMBINED TEXT STATS =====")
        print("chars:", len(combined))
        print("preview(first):")
        print(combined[:preview_len])
        return combined

    def _document_path_hint(self, doc: dict) -> str:
        return (
            (doc.get("name") or "").strip()
            or (doc.get("filename") or "").strip()
            or self._determine_document_filename(doc)
        )

    def _document_priority_score(self, doc: dict) -> int:
        path_hint = self._document_path_hint(doc)
        path_norm = f" {path_hint.lower().replace('\\', '/')} "
        basename = os.path.basename(path_norm.strip())
        score = 100

        if any(marker in path_norm for marker in self._TZ_STRONG_MARKERS):
            score -= 80
        if any(marker in path_norm for marker in self._TZ_WEAK_MARKERS):
            score -= 50
        if basename.startswith("тз") or basename.endswith("тз.docx") or basename.endswith("тз.doc"):
            score -= 40
        if "документация" in path_norm:
            score -= 20
        if "техническая часть" in path_norm:
            score -= 20
        if "приложение" in path_norm:
            score += 5
        if path_hint.lower().endswith(".pdf"):
            score += 5

        return score

    def _document_priority_sort_key(self, doc: dict):
        path_hint = self._document_path_hint(doc)
        score = self._document_priority_score(doc)
        return (
            score,
            len(path_hint),
            path_hint.lower(),
        )
        print("preview(last):")
        print(combined[-preview_len:] if len(combined) > preview_len else combined)

        print("================================\n")

        logger.info("Combined text chars=%s", len(combined))
        logger.debug("Combined text preview(first %s): %s", preview_len, combined[:preview_len])
        logger.debug("Combined text preview(last %s): %s", preview_len,
                     combined[-preview_len:] if len(combined) > preview_len else combined)

        return combined

    def get_application_deadline(self, lot_url: str) -> str | None:
        """
        Возвращает 'окончание подачи заявок' или None.
        Работает для tektorg и zakupki.
        """
        try:
            resp = self.session.get(lot_url, timeout=CONFIG["REQUEST_TIMEOUT"])
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "html.parser")

            # 1) tektorg: ищем label + <time>
            label = soup.find(
                string=lambda x: x and "оконч" in x.lower() and "подач" in x.lower() and "заяв" in x.lower())
            if label:
                parent = label.parent
                time_tag = parent.find_next("time") if parent else None
                if time_tag:
                    txt = time_tag.get_text(" ", strip=True)
                    return " ".join(txt.split()) if txt else None

            # 2) zakupki: ищем по заголовку блока, значение рядом
            for title_sel, value_sel in [
                (".common-text__title", ".common-text__value"),
                (".data-block__title", ".data-block__value"),
            ]:
                for t in soup.select(title_sel):
                    tt = t.get_text(" ", strip=True).lower()
                    if "оконч" in tt and "подач" in tt and "заяв" in tt:
                        v = t.find_next(class_=value_sel.replace(".", ""))
                        if v:
                            txt = v.get_text(" ", strip=True)
                            return " ".join(txt.split()) if txt else None

            # 3) fallback regex по всему тексту страницы
            text = soup.get_text(" ", strip=True)
            m = re.search(
                r"(?:окончание|дата)\s*(?:приема|подачи)\s*заявок.*?(\d{1,2}\.\d{1,2}\.\d{4})(?:\s*[,:]?\s*(\d{1,2}:\d{2}))?",
                text,
                flags=re.I,
            )
            if m:
                d = m.group(1)
                tm = m.group(2)
                return f"{d} {tm}".strip() if tm else d

            return None
        except Exception as e:
            logger.warning("get_application_deadline failed: %s", e)
            return None

    def _sanitize_for_llm(self, text: str, max_chars: int = 25_000) -> str:
        """
        Агрессивная очистка + ограничение размера.
        Убираем:
        - управляющие символы
        - мусорные последовательности
        - слишком длинные строки/повторы
        """
        if not text:
            return ""

        # нормализуем переносы и пробелы
        t = text.replace("\r\n", "\n").replace("\r", "\n")
        t = re.sub(r"[ \t\f\v]+", " ", t)

        # убираем control chars (кроме \n и \t)
        t = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", " ", t)

        # убираем “кракозябры” типа �
        t = t.replace("\ufffd", " ")

        # режем очень длинные “линии” (часто это мусор из PDF/сканов)
        t = "\n".join(ln[:2000] for ln in t.split("\n"))

        # схлопываем много одинаковых знаков подряд
        t = re.sub(r"([=*_#\-])\1{8,}", r"\1\1\1", t)

        # схлопываем множественные переносы
        t = re.sub(r"\n{3,}", "\n\n", t)

        t = t.strip()
        if not t:
            return ""

        # ограничиваем размер: берём начало + конец (обычно в начале ТЗ, в конце приложения)
        if len(t) > max_chars:
            head = t[: int(max_chars * 0.7)]
            tail = t[-int(max_chars * 0.3):]
            t = head + "\n\n[...ОБРЕЗАНО...]\n\n" + tail

        return t

    def call_llm_lot_analysis(self, combined_text: str) -> dict:
        """
        Запрос в LLM. Возвращает dict:
        { "goals_tasks": "...", "timelines": "...", "requirements": "..." }
        """
        provider = (CONFIG.get("LLM_PROVIDER") or "cloudru").strip().lower()
        providers = CONFIG.get("LLM_PROVIDERS") or {}
        provider_cfg = providers.get(provider, {})
        provider_key = (CONFIG.get("LLM_PROVIDER_KEYS") or {}).get(provider, "")

        env_key = provider_cfg.get("env_key") or ""
        api_key = (
            CONFIG.get("LLM_API_KEY")
            or provider_key
            or (os.environ.get(env_key) if env_key else "")
            or os.environ.get("API_KEY")
            or os.environ.get("MISTRAL_API_KEY")
            or ""
        ).strip()
        if not api_key:
            raise RuntimeError("LLM key not set. Set CONFIG['LLM_API_KEY'] or CONFIG['LLM_PROVIDER_KEYS'] or env var.")

        url = (
            CONFIG.get("LLM_API_URL")
            or provider_cfg.get("api_url")
            or "https://api.mistral.ai/v1/chat/completions"
        ).strip()
        model = (
            CONFIG.get("LLM_MODEL")
            or provider_cfg.get("model")
            or "mistral-large-latest"
        ).strip()

        # Если документов нет/текст пустой — сразу вернём прочерки, без вызова API
        if not combined_text:
            return {
                "subject": "—",
                "work_scope": "—",
                "work_and_submission_timelines": "Сроки работ: —; Сроки подачи: —",
                "fit_summary": "—",
                "final_verdict": "Неопределенно",
            }

        system_prompt = (
            "Ты помощник по анализу закупок для компании ООО «Современные Технологии».\n\n"

            "Контекст компании (важно для оценки соответствия):\n"
            "ООО «Современные Технологии» специализируется на управленческом, производственном и "
            "организационно-технологическом консалтинге, а также на сопровождении цифровых и "
            "организационных преобразований предприятий различных отраслей, включая промышленность "
            "и электроэнергетику.\n\n"

            "Компания выполняет:\n"
            "— управленческий и производственный консалтинг;\n"
            "— анализ, описание и оптимизацию бизнес-процессов;\n"
            "— обследование деятельности предприятий (в т.ч. эксплуатация, ТОиР, управление активами);\n"
            "— разработку методик, регламентов, моделей, показателей эффективности;\n"
            "— проекты цифровизации и внедрения прикладных систем как инструмента изменений;\n"
            "— внедрение, настройку, адаптацию и сопровождение ИТ-систем (EAM/ТОиР, аналитика, НСИ, витрины данных);\n"
            "— нормализацию и управление корпоративными данными и нормативно-технической информацией;\n"
            "— разработку отчетов, форм, моделей данных, интеграций;\n"
            "— low-code / no-code доработки и конфигурации (без сложной программной разработки ядра систем).\n\n"

            "Важно:\n"
            "Компания НЕ занимается физическим выполнением работ (ремонт, монтаж, ПНР «руками») и не "
            "выполняет сложную заказную разработку программного обеспечения «с нуля».\n"
            "ИТ-решения рассматриваются как инструмент управленческих и организационных изменений.\n\n"

            "Не является профильным ТОЛЬКО если это основной предмет закупки:\n"
            "— чистая поставка товаров, материалов, оборудования;\n"
            "— поставка лицензий или ПО без услуг по внедрению, настройке, обследованию или сопровождению.\n\n"

            "Если по тексту явно видно, что предмет закупки — исключительно поставка товаров/материалов "
            "(признаки: «поставка», «товар», «ТТН», «упаковка», «маркировка», «сертификат на продукцию», "
            "«ГОСТ на продукцию») — верни строго:\n"
            "{\"subject\":\"—\",\"work_scope\":\"—\",\"work_and_submission_timelines\":\"—\","
            "\"fit_summary\":\"Не наш профиль: поставка товаров/материалов без услуг.\","
            "\"final_verdict\":\"Не подходит\"}\n\n"

            "Тебе дан объединённый текст документов по одному лоту.\n"
            "Запрещено додумывать: используй только то, что явно есть в тексте.\n"
            "Если информации нет — ставь ровно символ \"—\".\n\n"

            "Верни ТОЛЬКО валидный JSON (без markdown, без пояснений), строго с ключами:\n"
            "{\"subject\":\"...\",\"work_scope\":\"...\",\"work_and_submission_timelines\":\"...\","
            "\"fit_summary\":\"...\",\"final_verdict\":\"...\"}\n\n"

            "Что нужно извлечь:\n\n"

            "1) subject — предмет закупки (1–2 предложения):\n"
            "Что требуется сделать / оказать / разработать и какой ожидается результат.\n"
            "Не включай цену и НМЦК.\n\n"

            "2) work_scope — состав работ / услуг:\n"
            "Перечисли все явно указанные работы и результаты через \"; \".\n"
            "Включай обследование, анализ, методики, регламенты, проектирование, внедрение, настройку, "
            "сопровождение, отчеты, формы, модели, интеграции, опытную эксплуатацию.\n"
            "Не включай финансовые условия и оплату.\n\n"

            "3) work_and_submission_timelines — сроки работ и сроки подачи:\n"
            "Одна строка, два блока:\n"
            "— \"Сроки работ: ...\";\n"
            "— \"Сроки подачи: ...\".\n"
            "Включай только явно указанные сроки.\n"
            "Если сроков нет — ставь \"—\" для соответствующего блока.\n\n"

            "4) fit_summary — оценка соответствия профилю ООО «Современные Технологии»:\n"
            "2–5 предложений.\n"
            "Обязательно:\n"
            "— укажи, почему лот подходит / потенциально подходит / не подходит;\n"
            "— если есть хотя бы косвенные признаки релевантности (консалтинг, анализ, внедрение, методики, "
            "сопровождение, ТОиР, EAM, бизнес-процессы, нормализация данных) — НЕ бракуй жёстко;\n"
            "— если информации недостаточно для уверенного вывода — прямо укажи это.\n\n"

            "5) final_verdict — итоговый вердикт:\n"
            "Разрешены ТОЛЬКО значения:\n"
            "— \"Подходит\";\n"
            "— \"Не подходит\";\n"
            "— \"Неопределенно\".\n\n"

            "Правило выбора вердикта:\n"
            "— \"Подходит\" — если лот явно связан с консалтингом, анализом, внедрением, методологией, сопровождением.\n"
            "— \"Неопределенно\" — если формулировки общие или недостаточно детализации, но есть признаки потенциальной релевантности.\n"
            "— \"Не подходит\" — только при явном несоответствии профилю (чистая поставка или физические работы).\n"
        )

        safe_text = self._sanitize_for_llm(combined_text, max_chars=25_000)

        print(f"[LLM INPUT] raw={len(combined_text)} chars, sanitized={len(safe_text)} chars")
        logger.info("[LLM INPUT] raw=%s chars, sanitized=%s chars", len(combined_text), len(safe_text))

        # ✅ ВЫВОД ПОЛНОГО ТЕКСТА, КОТОРЫЙ УЙДЕТ В МОДЕЛЬ
        print("\n===== LLM INPUT TEXT BEGIN =====\n")
        print(safe_text)
        print("\n===== LLM INPUT TEXT END =====\n")

        if not safe_text:
            return {
                "subject": "—",
                "work_scope": "—",
                "work_and_submission_timelines": "Сроки работ: —; Сроки подачи: —",
                "fit_summary": "—",
                "final_verdict": "Неопределенно",
            }

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": safe_text},
            ],
            "temperature": 0.2,
            "max_tokens": 4000,
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "application/json",
            # КРИТИЧНО: не просим br, чтобы requests мог распаковать ответ
            "Accept-Encoding": "gzip, deflate",
        }

        try:
            resp = self.session.post(
                url,
                headers=headers,
                json=payload,
                timeout=CONFIG.get("LLM_REQUEST_TIMEOUT", CONFIG["REQUEST_TIMEOUT"]),
            )
            resp.raise_for_status()
        except requests.RequestException as e:
            logger.warning("LLM request failed once: %s. Retrying...", e)
            time.sleep(1.0)
            resp = self.session.post(
                url,
                headers=headers,
                json=payload,
                timeout=CONFIG.get("LLM_REQUEST_TIMEOUT", CONFIG["REQUEST_TIMEOUT"]),
            )
            resp.raise_for_status()

        print("LLM HTTP status:", resp.status_code)
        print("LLM content-type:", resp.headers.get("content-type"))
        print("LLM content-encoding:", resp.headers.get("content-encoding"))
        print("LLM first 200 bytes:", (resp.content or b"")[:200])
        # 1) Безопасно читаем JSON ответа
        try:
            resp_json = resp.json()
        except Exception as err:
            print("\n===== LLM RESPONSE JSON ERROR =====")
            print("Cannot parse response as JSON:", err)
            print("Raw resp.text(first 2000):", (resp.text or "")[:2000])
            print("===================================\n")
            logger.warning("LLM response is not JSON: %s; text=%s", err, (resp.text or "")[:2000])
            return {
                "subject": "—",
                "work_scope": "—",
                "work_and_submission_timelines": "Сроки работ: —; Сроки подачи: —",
                "fit_summary": "—",
                "final_verdict": "Неопределенно",
            }

        # 2) Достаём content максимально аккуратно
        content = ""
        try:
            content = (resp_json.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        except Exception:
            content = ""

        print("\n===== LLM RAW RESPONSE =====\n", content, "\n============================\n")
        logger.info("LLM raw response len=%s", len(content))

        # 3) Извлекаем JSON-объект из ответа (даже если вокруг текст/```json)
        def _extract_json_object(text: str) -> dict:
            if not text:
                return {}

            cleaned = text.strip()
            cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r"\s*```$", "", cleaned)

            l = cleaned.find("{")
            r = cleaned.rfind("}")
            if l == -1 or r == -1 or r <= l:
                return {}

            candidate = cleaned[l:r + 1].strip()
            if not candidate:
                return {}

            try:
                return json.loads(candidate)
            except Exception as parse_err:
                print("===== LLM JSON PARSE ERROR =====")
                print(str(parse_err))
                print("Candidate(first 2000):", candidate[:2000])
                print("================================")
                logger.warning("LLM JSON parse error: %s; candidate=%s", parse_err, candidate[:2000])
                return {}

        data = _extract_json_object(content)

        print("\n===== LLM PARSED JSON =====\n", json.dumps(data, ensure_ascii=False, indent=2),
              "\n===========================\n")
        logger.info("LLM parsed json: %s", json.dumps(data, ensure_ascii=False))

        def _norm(v: str) -> str:
            v = (v or "").strip()
            return v if v else "—"

        result = {
            "subject": _norm(data.get("subject")),
            "work_scope": _norm(data.get("work_scope")),
            "work_and_submission_timelines": _norm(data.get("work_and_submission_timelines")),
            "fit_summary": _norm(data.get("fit_summary")),
            "final_verdict": _norm(data.get("final_verdict")),
        }

        # страховка: если модель вернула просто '—'
        if result["work_and_submission_timelines"] == "—":
            result["work_and_submission_timelines"] = "Сроки работ: —; Сроки подачи: —"
        if result["final_verdict"] not in ("Подходит", "Не подходит", "Неопределенно"):
            result["final_verdict"] = "Неопределенно"

        return result

    def _is_tektorg_allowed_doc_url(self, url: str) -> bool:
        try:
            p = urlparse(url)
            host = (p.netloc or "").lower()
            path = (p.path or "")

            # 1) Файлы лота
            if host == "44.tektorg.ru" and path.startswith("/file/get/"):
                return True

            # 2) Open API документы процедуры
            if host == "api.tektorg.ru" and path.startswith("/open-api/documents/procedure/"):
                return True

            return False
        except Exception:
            return False


    def _is_blacklisted_document_url(self, url: str) -> tuple[bool, str]:
        """
        Единый denylist для всех источников.
        Возвращает (True, reason) если URL нужно пропустить.
        """
        try:
            p = urlparse(url)
            host = (p.netloc or "").lower()
            path = (p.path or "").lower()

            # tektorg: режем страницы сайта (это не документы лота)
            if host in ("www.tektorg.ru", "tektorg.ru"):
                return True, "tektorg:www-page"

            # tektorg: режем страницы ".../documents" (это HTML-раздел, не файл)
            if host.endswith("tektorg.ru") and path.endswith("/documents"):
                return True, "tektorg:documents-page"

            # zakupki.gov.ru: не нужны *view.html
            if "zakupki.gov.ru" in host and path.endswith("view.html"):
                return True, "zakupki:view.html"

            if "zakupki.gov.ru" in host and path.endswith("documents.html"):
                return True, "zakupki:documents.html"

            if "zakupki.gov.ru" in host and path.endswith("zakupki-traffic.xlsx"):
                return True, "zakupki:zakupki-traffic.xlsx"

            # zakupki.gov.ru: служебные страницы
            if "zakupki.gov.ru" in host and "/purchase/public/download/signs/" in path:
                return True, "zakupki:signs-render"
            if "zakupki.gov.ru" in host and "/purchase/public/print-form/" in path:
                return True, "zakupki:print-form"

            if host.endswith("tektorg.ru") and path.endswith("/documents"):
                return True, "tektorg:documents-page"

            if host.endswith("tektorg.ru") and path.startswith("/documents/"):
                return True, "tektorg:site-documents"

            return False, ""
        except Exception:
            return False, ""

    def _looks_like_garbage_text(self, text: str) -> bool:
        """
        Простая эвристика: если почти нет букв (кириллица/латиница),
        то это похоже на бинарный мусор после decode().
        """
        if not text:
            return True

        t = text.strip()
        if len(t) < 50:
            return True

        letters = sum(1 for ch in t if ch.isalpha())
        ratio = letters / max(len(t), 1)

        # если букв меньше ~12% — это почти наверняка мусор
        return ratio < 0.12

    def _is_archive_document(self, filename: str, content_type: str = "") -> bool:
        name = (filename or "").lower()
        ctype = (content_type or "").lower()
        if any(name.endswith(ext) for ext in self._ARCHIVE_EXTENSIONS):
            return True
        return "zip" in ctype

    def _is_extractable_document(self, filename: str) -> bool:
        name = (filename or "").lower()
        return any(name.endswith(ext) for ext in self._TEXT_EXTRACT_EXTENSIONS)

    def _expand_downloaded_documents(
        self,
        documents: list[dict],
        progress_callback=None,
    ) -> list[dict]:
        expanded: list[dict] = []
        total_archives = sum(
            1
            for doc in documents
            if self._is_archive_document(
                self._determine_document_filename(doc),
                doc.get("content_type", ""),
            )
        )
        archive_idx = 0

        for doc in documents:
            filename = self._determine_document_filename(doc)
            if not self._is_archive_document(filename, doc.get("content_type", "")):
                expanded.append(doc)
                continue

            archive_idx += 1
            if progress_callback:
                try:
                    progress_callback(
                        f"Распаковка архива {archive_idx}/{max(total_archives, 1)}: {doc.get('name') or filename}"
                    )
                except Exception:
                    pass

            extracted = self._extract_documents_from_archive(doc, progress_callback)
            if extracted:
                expanded.extend(extracted)
            else:
                logger.info(
                    "Archive kept without extraction results: %s",
                    doc.get("name") or filename,
                )
                expanded.append(doc)

        return expanded

    def _extract_documents_from_archive(
        self,
        doc: dict,
        progress_callback=None,
        depth: int = 0,
        parent_path: str | None = None,
    ) -> list[dict]:
        filename = self._determine_document_filename(doc)
        if filename.lower().endswith(".zip") or "zip" in (doc.get("content_type", "") or "").lower():
            return self._extract_documents_from_zip(
                doc,
                depth=depth,
                progress_callback=progress_callback,
                parent_path=parent_path,
            )
        if filename.lower().endswith((".rar", ".7z")):
            return self._extract_documents_with_7z(
                doc,
                depth=depth,
                progress_callback=progress_callback,
                parent_path=parent_path,
            )
        return []

    def _extract_documents_from_zip(
        self,
        doc: dict,
        depth: int = 0,
        progress_callback=None,
        parent_path: str | None = None,
    ) -> list[dict]:
        if depth > self._ARCHIVE_MAX_DEPTH:
            logger.info("Archive nesting depth exceeded for %s", doc.get("name"))
            return []

        archive_name = parent_path or (doc.get("name") or self._determine_document_filename(doc))
        extracted_docs: list[dict] = []
        total_size = 0
        logger.info("Открытие ZIP-архива для распаковки: %s", archive_name)

        try:
            with zipfile.ZipFile(BytesIO(doc.get("content") or b"")) as zf:
                members = [m for m in zf.infolist() if not m.is_dir()]
                if len(members) > self._ARCHIVE_MAX_FILES:
                    logger.warning(
                        "Archive %s has too many files (%s), truncating to %s",
                        archive_name,
                        len(members),
                        self._ARCHIVE_MAX_FILES,
                    )
                    members = members[: self._ARCHIVE_MAX_FILES]

                for member in members:
                    inner_name = (member.filename or "").replace("\\", "/").strip("/")
                    if not inner_name:
                        continue

                    basename = os.path.basename(inner_name)
                    if basename.startswith(".") or basename.startswith("~$"):
                        logger.info(
                            "Пропуск скрытого/временного файла в архиве %s: %s",
                            archive_name,
                            inner_name,
                        )
                        continue

                    total_size += max(member.file_size, 0)
                    if total_size > self._ARCHIVE_MAX_TOTAL_SIZE:
                        logger.warning(
                            "Archive %s exceeds size limit after %s bytes",
                            archive_name,
                            total_size,
                        )
                        break

                    try:
                        content = zf.read(member)
                    except RuntimeError as e:
                        logger.warning("Failed to read zip member %s: %s", inner_name, e)
                        continue
                    except Exception as e:
                        logger.warning("Unexpected zip read error for %s: %s", inner_name, e)
                        continue

                    composed_name = f"{archive_name} -> {inner_name}"
                    content_type = self._guess_content_type_by_name(inner_name)
                    logger.info(
                        "Найден файл в архиве: архив=%r файл=%r размер=%s",
                        archive_name,
                        inner_name,
                        len(content),
                    )
                    nested_doc = {
                        "name": composed_name,
                        "filename": inner_name,
                        "content": content,
                        "size": len(content),
                        "url": doc.get("url"),
                        "content_type": content_type,
                        "source_archive": archive_name,
                    }

                    if self._is_archive_document(inner_name, content_type):
                        logger.info(
                            "Найден вложенный архив: архив=%r файл=%r",
                            archive_name,
                            inner_name,
                        )
                        nested = self._extract_documents_from_archive(
                            nested_doc,
                            progress_callback=progress_callback,
                            depth=depth + 1,
                            parent_path=composed_name,
                        )
                        if nested:
                            extracted_docs.extend(nested)
                            continue

                    if self._is_extractable_document(inner_name):
                        logger.info(
                            "Файл из архива принят в анализ: архив=%r файл=%r",
                            archive_name,
                            inner_name,
                        )
                        extracted_docs.append(nested_doc)
                    else:
                        logger.info(
                            "Файл из архива пропущен по расширению: архив=%r файл=%r",
                            archive_name,
                            inner_name,
                        )

        except zipfile.BadZipFile as e:
            logger.warning("Bad ZIP archive %s: %s", archive_name, e)
            return []
        except Exception as e:
            logger.warning("Archive extraction failed for %s: %s", archive_name, e)
            return []

        return extracted_docs

    def _find_7z_executable(self) -> str | None:
        candidates = [
            r"C:\Program Files\7-Zip\7z.exe",
            r"C:\Program Files (x86)\7-Zip\7z.exe",
        ]
        for path in candidates:
            if os.path.exists(path):
                return path
        return shutil.which("7z")

    def _extract_documents_with_7z(
        self,
        doc: dict,
        depth: int = 0,
        progress_callback=None,
        parent_path: str | None = None,
    ) -> list[dict]:
        if depth > self._ARCHIVE_MAX_DEPTH:
            logger.info("Archive nesting depth exceeded for %s", doc.get("name"))
            return []

        seven_zip = self._find_7z_executable()
        if not seven_zip:
            logger.warning("7z.exe not found, cannot extract archive %s", doc.get("name"))
            return []

        archive_name = parent_path or (doc.get("name") or self._determine_document_filename(doc))
        temp_archive_path = None
        temp_dir = None
        extracted_docs: list[dict] = []
        logger.info("Открытие архива через 7z для распаковки: %s", archive_name)

        try:
            suffix = os.path.splitext(self._determine_document_filename(doc))[1] or ".bin"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_archive:
                tmp_archive.write(doc.get("content") or b"")
                temp_archive_path = tmp_archive.name

            temp_dir = tempfile.mkdtemp(prefix="lot_archive_")
            cmd = [seven_zip, "x", "-y", f"-o{temp_dir}", temp_archive_path]
            completed = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
            if completed.returncode not in (0, 1):
                logger.warning(
                    "7z extraction failed for %s: %s",
                    archive_name,
                    (completed.stderr or completed.stdout or "").strip(),
                )
                return []

            total_size = 0
            total_files = 0
            for root, _dirs, files in os.walk(temp_dir):
                files.sort()
                for file_name in files:
                    total_files += 1
                    if total_files > self._ARCHIVE_MAX_FILES:
                        logger.warning(
                            "Archive %s has too many files (%s), truncating",
                            archive_name,
                            total_files,
                        )
                        return extracted_docs

                    full_path = os.path.join(root, file_name)
                    rel_path = os.path.relpath(full_path, temp_dir).replace("\\", "/")
                    if file_name.startswith(".") or file_name.startswith("~$"):
                        logger.info(
                            "Пропуск скрытого/временного файла в архиве %s: %s",
                            archive_name,
                            rel_path,
                        )
                        continue

                    try:
                        with open(full_path, "rb") as f:
                            content = f.read()
                    except Exception as e:
                        logger.warning("Failed to read extracted file %s: %s", rel_path, e)
                        continue

                    total_size += len(content)
                    if total_size > self._ARCHIVE_MAX_TOTAL_SIZE:
                        logger.warning(
                            "Archive %s exceeds size limit after %s bytes",
                            archive_name,
                            total_size,
                        )
                        return extracted_docs

                    composed_name = f"{archive_name} -> {rel_path}"
                    content_type = self._guess_content_type_by_name(rel_path)
                    logger.info(
                        "Найден файл в архиве: архив=%r файл=%r размер=%s",
                        archive_name,
                        rel_path,
                        len(content),
                    )
                    nested_doc = {
                        "name": composed_name,
                        "filename": rel_path,
                        "content": content,
                        "size": len(content),
                        "url": doc.get("url"),
                        "content_type": content_type,
                        "source_archive": archive_name,
                    }

                    if self._is_archive_document(rel_path, content_type):
                        logger.info(
                            "Найден вложенный архив: архив=%r файл=%r",
                            archive_name,
                            rel_path,
                        )
                        nested = self._extract_documents_from_archive(
                            nested_doc,
                            progress_callback=progress_callback,
                            depth=depth + 1,
                            parent_path=composed_name,
                        )
                        if nested:
                            extracted_docs.extend(nested)
                            continue

                    if self._is_extractable_document(rel_path):
                        logger.info(
                            "Файл из архива принят в анализ: архив=%r файл=%r",
                            archive_name,
                            rel_path,
                        )
                        extracted_docs.append(nested_doc)
                    else:
                        logger.info(
                            "Файл из архива пропущен по расширению: архив=%r файл=%r",
                            archive_name,
                            rel_path,
                        )

        except Exception as e:
            logger.warning("7z extraction failed for %s: %s", archive_name, e)
            return []
        finally:
            try:
                if temp_archive_path and os.path.exists(temp_archive_path):
                    os.remove(temp_archive_path)
            except Exception:
                pass
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass

        return extracted_docs

    def _guess_content_type_by_name(self, filename: str) -> str:
        name = (filename or "").lower()
        mapping = {
            ".pdf": "application/pdf",
            ".doc": "application/msword",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xls": "application/vnd.ms-excel",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
            ".txt": "text/plain",
            ".rtf": "application/rtf",
            ".zip": "application/zip",
            ".rar": "application/vnd.rar",
            ".7z": "application/x-7z-compressed",
            ".odt": "application/vnd.oasis.opendocument.text",
            ".ods": "application/vnd.oasis.opendocument.spreadsheet",
        }
        for ext, ctype in mapping.items():
            if name.endswith(ext):
                return ctype
        return "application/octet-stream"

    def __init__(self, sources: list[ProcurementSource] | None = None):
        self.session = self._create_session()
        # По умолчанию используем все доступные источники
        if sources is None:
            self.sources = [ZakupkiGovSource(), TektorgSource(), BidzaarSource()]
        else:
            self.sources = sources

    # создание сессии
    def _create_session(self):
        """Create a requests session with retry strategy and proper headers."""
        session = requests.Session()

        # Setup retry strategy
        retry_strategy = Retry(
            total=CONFIG["MAX_RETRIES"],
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)

        # Set headers
        session.headers.update(
            {
                "User-Agent": CONFIG["USER_AGENT"],
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                "Accept-Language": "ru-RU,ru;q=0.9,en;q=0.8",
                "Accept-Encoding": "gzip, deflate",
                "Connection": "keep-alive",
                "Upgrade-Insecure-Requests": "1",
                "Sec-Fetch-Dest": "document",
                "Sec-Fetch-Mode": "navigate",
                "Sec-Fetch-Site": "none",
            }
        )

        return session

    def search_procurements(
            self,
            keyword,
            min_price=None,
            max_price=None,
            purchase_stage=None,
            law=None,
            progress_callback=None,
            source_names: list[str] | None = None,
    ):
        """
        Search for procurements across multiple sources.
        
        Args:
            source_names: список названий источников для поиска (если None - все источники)
        """
        if not keyword or not keyword.strip():
            raise ValueError("Ключевое слово не может быть пустым")

        # Фильтруем источники, если указаны
        sources_to_use = self.sources
        if source_names:
            sources_to_use = [s for s in self.sources if s.get_name() in source_names]

        if not sources_to_use:
            raise ValueError("Не выбрано ни одного источника для поиска")

        all_results = []

        def fetch_page(source, source_name, page_number):
            url, params = source.build_search_url(
                keyword, min_price, max_price, purchase_stage, law, page_number=page_number
            )

            logger.info(
                f"Searching on {source_name}: {keyword!r} page={page_number} with filters - "
                f"stage: {purchase_stage}, law: {law}, "
                f"min_price: {min_price}, max_price: {max_price}"
            )

            logger.info(f"Requesting {url} with params: {params}")

            # Для tektorg.ru используем специальную обработку параметров
            if source_name == "tektorg.ru":
                from urllib.parse import quote_plus, quote
                query_parts = []
                for key, value in params.items():
                    if value:
                        encoded_key = quote(key, safe="")
                        if isinstance(value, str):
                            if "+" in value and key in ("sumPrice_start", "sumPrice_end"):
                                encoded_value = quote(value, safe="+")
                            else:
                                encoded_value = quote_plus(value)
                        else:
                            encoded_value = quote_plus(str(value))
                        query_parts.append(f"{encoded_key}={encoded_value}")

                if query_parts:
                    full_url = f"{url}?{'&'.join(query_parts)}"
                else:
                    full_url = url

                logger.info(f"Tektorg final URL: {full_url}")
                if progress_callback:
                    progress_callback(f"Запрос: {full_url}")
                response = self.session.get(
                    full_url, timeout=CONFIG["REQUEST_TIMEOUT"]
                )
            else:
                response = self.session.get(
                    url, params=params, timeout=CONFIG["REQUEST_TIMEOUT"]
                )

            response.raise_for_status()

            logger.info(f"Final URL (after redirects): {response.url}")
            if progress_callback and source_name != "tektorg.ru":
                progress_callback(f"Запрос: {response.url}")

            if progress_callback:
                progress_callback(f"Обработка результатов с {source_name}...")

            time.sleep(CONFIG["REQUEST_DELAY"])

            html_content = response.text
            if not isinstance(html_content, str):
                html_content = response.content.decode('utf-8', errors='ignore')

            source_results = source.parse_results(html_content, progress_callback)

            filtered_results = []
            for result in source_results:
                if self._passes_price_filter(result.get("Цена", 0), min_price, max_price):
                    result["Источник"] = source_name
                    filtered_results.append(result)

            return filtered_results

        for source in sources_to_use:
            source_name = source.get_name()
            try:
                if progress_callback:
                    progress_callback(f"Поиск на {source_name}...")

                page = 1
                while True:
                    if progress_callback:
                        progress_callback(f"{source_name}: страница {page}")
                    page_results = fetch_page(source, source_name, page)
                    if not page_results:
                        break
                    all_results.extend(page_results)
                    logger.info(f"Found {len(page_results)} results from {source_name} page {page}")
                    page += 1

            except Exception as e:
                logger.error(f"Error searching on {source_name}: {e}")
                if progress_callback:
                    progress_callback(f"Ошибка на {source_name}: {str(e)}")
                continue

        logger.info(f"Total results from all sources: {len(all_results)}")
        return all_results

    def _normalize_title(self, title: str) -> str:
        t = (title or "").strip().lower()
        t = re.sub(r"\s+", " ", t)
        return t

    def _normalize_price_for_key(self, price) -> str:
        if price is None:
            return ""
        if isinstance(price, (int, float)):
            if price <= 0:
                return ""
            return str(int(round(price)))
        s = str(price).replace("\xa0", " ").strip()
        m = re.search(r"\d[\d\s.,]*\d|\d", s)
        if not m:
            return ""
        num = m.group(0).replace(" ", "").replace(",", ".")
        try:
            val = float(num)
        except Exception:
            return ""
        if val <= 0:
            return ""
        return str(int(round(val)))

    def make_lot_cache_key(self, title, price) -> str:
        return f"{self._normalize_title(title)}|{self._normalize_price_for_key(price)}"

    def call_llm_lot_title_filter(self, items: list[dict]) -> dict:
        if not items:
            return {}

        provider = (CONFIG.get("LLM_PROVIDER") or "cloudru").strip().lower()
        providers = CONFIG.get("LLM_PROVIDERS") or {}
        provider_cfg = providers.get(provider, {})
        provider_key = (CONFIG.get("LLM_PROVIDER_KEYS") or {}).get(provider, "")

        env_key = provider_cfg.get("env_key") or ""
        api_key = (
            CONFIG.get("LLM_API_KEY")
            or provider_key
            or (os.environ.get(env_key) if env_key else "")
            or os.environ.get("API_KEY")
            or os.environ.get("MISTRAL_API_KEY")
            or ""
        ).strip()
        if not api_key:
            raise RuntimeError("LLM key not set. Set CONFIG['LLM_API_KEY'] or CONFIG['LLM_PROVIDER_KEYS'] or env var.")

        url = (
            CONFIG.get("LLM_API_URL")
            or provider_cfg.get("api_url")
            or "https://api.mistral.ai/v1/chat/completions"
        ).strip()
        model = (
            CONFIG.get("LLM_MODEL")
            or provider_cfg.get("model")
            or "mistral-large-latest"
        ).strip()

        def _format_price(price) -> str:
            if price is None:
                return "не указана"
            if isinstance(price, (int, float)):
                if price <= 0:
                    return "не указана"
                return f"{price:,.2f}".replace(",", " ")
            s = str(price).strip()
            return s if s else "не указана"

        system_prompt = (
            "Ты помощник по анализу закупочных лотов для консалтинговой компании ООО «Современные Технологии».\n\n"

            "ОПИСАНИЕ КОМПАНИИ:\n"
            "ООО «Современные Технологии» — консалтинговая компания, специализирующаяся на комплексном внедрении изменений и "
            "цифровых решений для повышения эффективности деятельности предприятий различных отраслей.\n\n"

            "Компания реализует проекты, направленные на:\n"
            "— цифровую трансформацию бизнеса и производственных систем;\n"
            "— реинжиниринг и оптимизацию бизнес-процессов;\n"
            "— внедрение управленческих и производственных контуров (BPM, EAM, MES, MDM и аналогичных систем);\n"
            "— управление производственными активами и данными;\n"
            "— нормализацию, структурирование и использование нормативно-технической документации;\n"
            "— сопровождение организационных и управленческих изменений.\n\n"

            "Ключевая особенность компании:\n"
            "ИТ-решения рассматриваются исключительно как инструмент реализации управленческих и организационных изменений. "
            "Компания не занимается изолированной разработкой, поставкой или сопровождением ИТ-продуктов без консалтинговой "
            "и трансформационной составляющей.\n\n"

            "ЧТО ЯВЛЯЕТСЯ ПРОФИЛЕМ КОМПАНИИ:\n"
            "— управленческий и производственный консалтинг;\n"
            "— проекты по цифровой трансформации предприятий;\n"
            "— анализ, проектирование и внедрение целевых бизнес-процессов;\n"
            "— внедрение систем управления как части трансформационных программ;\n"
            "— работа с корпоративными данными и нормативно-технической информацией.\n\n"

            "ЧТО НЕ ЯВЛЯЕТСЯ ПРОФИЛЕМ КОМПАНИИ (ЕСЛИ ЛОТ ПРО ЭТО — СТАВЬ 1):\n"
            "— информационная безопасность в любом виде;\n"
            "— предоставление доступов, учетных записей, прав пользователей;\n"
            "— ИТ-инфраструктура, сети, СКС, серверы;\n"
            "— техническая поддержка, сопровождение и аутсорсинг;\n"
            "— поставка программного обеспечения, лицензий или оборудования без консалтинга;\n"
            "— обучение, типовые ИТ-услуги и сервисные контракты.\n\n"

            "ЗАДАЧА:\n"
            "На основе НАЗВАНИЯ лота (и описания, если оно есть) определить степень соответствия лота профилю компании.\n"
            "Цену НЕ учитывать.\n\n"

            "ШКАЛА ОЦЕНКИ:\n"
            "5 — прямое попадание в профиль компании, типовой консалтинговый или трансформационный проект;\n"
            "4 — близко к профилю, возможно участие после уточнений;\n"
            "3 — косвенное соответствие, интерес ограниченный;\n"
            "2 — слабое соответствие или высокая неопределённость;\n"
            "1 — не соответствует профилю компании.\n\n"

            "ПРАВИЛА ОЦЕНКИ:\n"
            "— если лот относится к НЕ профильным направлениям, ставь 1 без колебаний;\n"
            "— если в лоте отсутствуют явные признаки консалтинга, трансформации или управления изменениями — не ставь выше 3;\n"
            "— если формулировка общая, бюрократическая или неясная — понижай оценку;\n"
            "— сомнения всегда трактуй в сторону понижения оценки.\n\n"

            "ФОРМАТ ОТВЕТА:\n"
            "Верни строго JSON-объект, где ключ — номер лота, значение — целое число от 1 до 5.\n"
            "Никаких пояснений, только JSON."
        )

        lines = ["Список лотов:"]
        for item in items:
            idx = str(item.get("id", "")).strip()
            title = str(item.get("title", "")).strip()
            price = _format_price(item.get("price"))
            lines.append(f"{idx}) Название: {title}")
            lines.append(f"   Цена: {price}")
        user_prompt = "\n".join(lines)
        safe_text = self._sanitize_for_llm(user_prompt, max_chars=8000)

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": safe_text},
            ],
            "temperature": 0.0,
            "max_tokens": 1000,
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Accept-Encoding": "gzip, deflate",
        }

        try:
            resp = self.session.post(
                url,
                headers=headers,
                json=payload,
                timeout=CONFIG.get("LLM_REQUEST_TIMEOUT", CONFIG["REQUEST_TIMEOUT"]),
            )
            resp.raise_for_status()
        except requests.RequestException as e:
            logger.warning("LLM request failed once: %s. Retrying...", e)
            time.sleep(1.0)
            resp = self.session.post(
                url,
                headers=headers,
                json=payload,
                timeout=CONFIG.get("LLM_REQUEST_TIMEOUT", CONFIG["REQUEST_TIMEOUT"]),
            )
            resp.raise_for_status()

        try:
            resp_json = resp.json()
        except Exception as err:
            logger.warning("LLM response is not JSON: %s; text=%s", err, (resp.text or "")[:2000])
            return {str(item.get("id")): "не подходит" for item in items}

        content = ""
        try:
            content = (resp_json.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        except Exception:
            content = ""

        try:
            debug_path = os.path.join(os.getcwd(), "llm_title_raw.txt")
            with open(debug_path, "a", encoding="utf-8") as f:
                f.write("\n===== LLM RAW RESPONSE =====\n")
                f.write(content or "")
                f.write("\n============================\n")
        except Exception as e:
            logger.warning("Failed to write LLM raw response: %s", e)

        def _extract_json_object(text: str) -> dict:
            if not text:
                return {}
            cleaned = text.strip()
            cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r"\s*```$", "", cleaned)
            l = cleaned.find("{")
            r = cleaned.rfind("}")
            if l == -1 or r == -1 or r <= l:
                return {}
            candidate = cleaned[l:r + 1].strip()
            if not candidate:
                return {}
            try:
                return json.loads(candidate)
            except Exception as parse_err:
                logger.warning("LLM JSON parse error: %s; candidate=%s", parse_err, candidate[:2000])
                return {}

        data = _extract_json_object(content)

        result = {}
        for item in items:
            key = str(item.get("id"))
            raw = data.get(key)
            score = 1
            try:
                score = int(raw)
            except Exception:
                try:
                    score = int(str(raw).strip())
                except Exception:
                    score = 1
            if score < 1 or score > 5:
                score = 1
            result[key] = score

        return result

    def _normalize_keywords(self, keywords):
        norm = []
        for kw in keywords or []:
            if kw is None:
                continue
            if not isinstance(kw, str):
                kw = str(kw)
            kw = kw.strip().lower()
            if kw:
                norm.append(kw)
        return norm

    def filter_lots_by_content(
        self,
        lots,
        include_keywords,
        blacklist=None,
        top_n=10,
        progress_callback=None,
    ):
        include = self._normalize_keywords(include_keywords)
        deny = self._normalize_keywords(blacklist)
        if not include and not deny:
            return lots
        to_check = lots[:top_n] if top_n else lots
        filtered = []
        total = len(to_check)

        for i, lot in enumerate(to_check, start=1):
            if progress_callback:
                progress_callback(f"Prefilter {i}/{total}: download + scan")

            lot_title = (lot.get("Название") or "").strip()
            lot_url = (lot.get("Ссылка") or "").strip()
            combined_text = lot_title

            documents = []
            if lot_url:
                try:
                    documents = self.download_documents(lot_url, progress_callback)
                except Exception as e:
                    logger.warning("Prefilter download failed for %s: %s", lot_url, e)

            if documents:
                combined_text = "\n".join(
                    [combined_text, self.build_lot_documents_text(documents)]
                )

            lower_text = combined_text.lower()
            if deny and any(bad in lower_text for bad in deny):
                continue

            if not include:
                filtered.append(lot)
                continue

            matches = [kw for kw in include if kw in lower_text]
            if matches:
                lot["_prefilter_matches"] = matches
                filtered.append(lot)

        return filtered

    def _get_stage_params(self, purchase_stage):
        """Get URL parameters for purchase stage filter."""
        stage_mapping = {
            "SUBMISSION": {"af": "on"},  # Подача заявок
            "EVALUATION": {"ca": "on"},  # Работа комиссии
            "COMPLETED": {"pc": "on"},  # Закупка завершена
            "CANCELLED": {"pa": "on"},  # Закупка отменена
        }
        return stage_mapping.get(purchase_stage, {})

    def _get_law_params(self, law):
        """Get URL parameters for law filter."""
        law_mapping = {
            "44FZ": {"fz44": "on"},
            "223FZ": {"fz223": "on"},
            "PP615": {"ppRf615": "on"},
        }
        return law_mapping.get(law, {})

    def _add_price_params(self, params: dict, min_price, max_price) -> None:
        """
        Добавляет в параметры запроса фильтр по цене так,
        как это делает сайт zakupki.gov.ru:

        priceFromGeneral — минимальная цена
        priceToGeneral   — максимальная цена
        currencyIdGeneral = -1 (любая валюта / по умолчанию)
        """

        def _to_int_or_none(value):
            if value is None:
                return None
            if isinstance(value, (int, float)):
                return int(value)

            # Строка из UI: убираем пробелы/неразрывные пробелы, меняем запятую на точку
            s = str(value).strip()
            if not s:
                return None
            s = s.replace("\u00a0", "").replace(" ", "")
            s = s.replace(",", ".")
            try:
                return int(float(s))
            except ValueError:
                logger.warning(f"Не удалось преобразовать цену '{value}' в число")
                return None

        min_val = _to_int_or_none(min_price)
        max_val = _to_int_or_none(max_price)

        if min_val is not None:
            params["priceFromGeneral"] = str(min_val)
        if max_val is not None:
            params["priceToGeneral"] = str(max_val)

        # как в примере запроса с сайта: -1 = любая валюта
        if "currencyIdGeneral" not in params:
            params["currencyIdGeneral"] = "-1"

        logger.info(
            "Фильтр по цене в запросе: priceFromGeneral=%s, priceToGeneral=%s, currencyIdGeneral=%s",
            params.get("priceFromGeneral"),
            params.get("priceToGeneral"),
            params.get("currencyIdGeneral"),
        )

    def _parse_results(
            self, html_content, min_price=None, max_price=None, progress_callback=None
    ):
        """Parse HTML content and extract procurement data."""
        try:
            soup = BeautifulSoup(html_content, "html.parser")

            # Multiple selector strategies for robustness
            selectors = [
                ".search-registry-entry-block",
                ".registry-entry",
                "[class*='registry-entry']",
                ".search-result-item",
            ]

            lots = []
            for selector in selectors:
                lots = soup.select(selector)
                if lots:
                    logger.info(f"Found {len(lots)} results using selector: {selector}")
                    break

            if not lots:
                logger.warning("No results found with any selector")
                return []

            results = []
            total_lots = len(lots)

            for i, lot in enumerate(lots):
                if progress_callback:
                    progress_callback(f"Обработка результата {i + 1} из {total_lots}...")

                try:
                    result = self._parse_single_lot(lot)
                    if result and self._passes_price_filter(
                            result["Цена"], min_price, max_price
                    ):
                        results.append(result)
                except Exception as e:
                    logger.warning(f"Failed to parse lot {i + 1}: {e}")
                    continue

            logger.info(f"Successfully parsed {len(results)} results")
            return results

        except Exception as e:
            logger.error(f"Failed to parse HTML: {e}")
            raise Exception("Ошибка обработки данных с сервера")

    def _passes_price_filter(self, price, min_price, max_price):
        """Check if price passes the min/max price filters."""
        if min_price is not None and price < min_price:
            return False
        if max_price is not None and price > max_price:
            return False
        return True

    def _parse_single_lot(self, lot):
        """Parse a single procurement lot with multiple fallback strategies."""
        title = self._extract_title(lot)
        price = self._extract_price(lot)
        link = self._extract_link(lot)

        if not title:
            return None

        return {"Название": title, "Цена": price, "Ссылка": link}

    def _extract_title(self, lot):
        """Extract title with multiple fallback selectors."""
        header_text = None
        object_text = None

        # 1. Заголовок сверху: "44-ФЗ Электронный аукцион"
        try:
            header_el = lot.select_one(".registry-entry__header-top__title")
            if header_el:
                header_text = header_el.get_text(" ", strip=True)
                if header_text:
                    header_text = " ".join(header_text.split())
        except Exception as e:
            logger.debug(f"Failed to extract header title: {e}")

        # 2. "Объект закупки" из body
        try:
            for body_block in lot.select(".registry-entry__body-block"):
                title_el = body_block.select_one(".registry-entry__body-title")
                if not title_el:
                    continue

                title_label = title_el.get_text(strip=True).lower()
                if "объект закупки" in title_label:
                    value_el = body_block.select_one(".registry-entry__body-value")
                    if value_el:
                        object_text = value_el.get_text(" ", strip=True)
                        if object_text:
                            object_text = " ".join(object_text.split())
                    break
        except Exception as e:
            logger.debug(f"Failed to extract object text: {e}")

        # 3. Комбинируем
        if header_text and object_text:
            return f"{header_text} — {object_text}"
        if object_text:
            return object_text
        if header_text:
            return header_text

        # 4. Фолбэк — старые селекторы на всякий случай
        title_selectors = [
            ".registry-entry__title",
            "[class*='title']",
            "h3",
            "h2",
            ".search-result-title",
        ]

        for selector in title_selectors:
            title_el = lot.select_one(selector)
            if title_el:
                title = title_el.get_text(" ", strip=True)
                if title and len(title) > 5:
                    return title

        return "Название не найдено"

    def _extract_price(self, lot):
        """Extract and validate price with multiple fallback selectors."""
        price_selectors = [
            ".price-block__value",
            ".price",
            "[class*='price']",
            ".cost",
            "[class*='cost']",
        ]

        for selector in price_selectors:
            price_el = lot.select_one(selector)
            if price_el:
                try:
                    price_text = (
                        price_el.get_text(strip=True)
                        .replace("\xa0", "")
                        .replace("₽", "")
                        .replace("руб", "")
                        .replace(" ", "")
                        .replace(",", ".")
                    )

                    numbers = re.findall(r"\d+\.?\d*", price_text)
                    if numbers:
                        price = float(numbers[0])
                        if price >= 0:
                            return price
                except (ValueError, AttributeError, IndexError) as e:
                    logger.debug(f"Price parsing failed for selector {selector}: {e}")
                    continue

        return 0.0

    def _extract_link(self, lot):
        base_url = CONFIG["BASE_URL"]

        link_el = lot.select_one(".registry-entry__header-mid__number a[href]")
        if link_el:
            href = link_el["href"].strip()
            if href:
                return urljoin(base_url, href)

        # 2. Альтернативный вариант — если структура чуть отличается
        candidates = lot.select("a[href*='notice/']")
        for link in candidates:
            href = (link.get("href") or "").strip()
            if not href:
                continue
            # пропускаем ссылки на printForm
            if "printForm" in href.lower():
                continue
            if "view" in href.lower() and "regNumber" in href:
                # ТУТ тоже нужно делать urljoin, иначе получаешь только хвост
                return urljoin(base_url, href)

        # 3. Если всё совсем плохо — явное "не найдено"
        return "Ссылка не найдена"


    # УЛУЧШЕННЫЕ МЕТОДЫ ДЛЯ РАБОТЫ С ДОКУМЕНТАМИ
    def _extract_notice_info_id(self, url):
        """Извлекает noticeInfoId из URL."""
        try:
            parsed_url = urlparse(url)
            query_params = parse_qs(parsed_url.query)
            notice_info_id = query_params.get("noticeInfoId", [None])[0]
            return notice_info_id
        except Exception as e:
            logger.error(f"Error extracting noticeInfoId from {url}: {e}")
            return None

    def _get_source_for_url(self, url: str) -> ProcurementSource | None:
        """Определяет источник по URL."""
        for source in self.sources:
            if source.get_name() in url:
                return source
        return None

    def _get_documents_url_legacy(self, lot_url):
        """Старая логика для zakupki.gov.ru (для обратной совместимости)."""
        parsed_url = urlparse(lot_url)
        path = parsed_url.path or ""
        query_dict = parse_qs(parsed_url.query)

        documents_url = None
        params = {}

        if "/epz/order/notice/" in path and "view/" in path:
            if "common-info.html" in path:
                docs_path = path.replace("common-info.html", "documents.html")
            else:
                docs_path = path
            documents_url = urljoin(CONFIG["BASE_URL"], docs_path)
            params = {k: v[0] for k, v in query_dict.items() if v}
        else:
            notice_info_id = query_dict.get("noticeInfoId", [None])[0]
            if not notice_info_id:
                notice_info_id = self._extract_notice_info_id(lot_url)
            if notice_info_id:
                documents_url = f"{CONFIG['BASE_URL']}/epz/order/notice/notice223/documents.html"
                params = {
                    "noticeInfoId": notice_info_id,
                    "backUrl": "/epz/order/notice/notice223/search.html",
                }

        return documents_url, params

    def download_documents(self, lot_url, progress_callback=None):
        """
        Скачивает документы из лота. Определяет источник по URL и использует соответствующий метод.
        """
        try:
            # Определяем источник по URL
            source = self._get_source_for_url(lot_url)
            source_name = source.get_name() if source else ""
            if source and hasattr(source, 'get_documents_url'):
                docs_info = source.get_documents_url(lot_url)
                if docs_info:
                    documents_url, params = docs_info
                else:
                    # Если источник не поддерживает get_documents_url, используем старую логику
                    documents_url, params = self._get_documents_url_legacy(lot_url)
            else:
                # Fallback на старую логику для zakupki.gov.ru
                documents_url, params = self._get_documents_url_legacy(lot_url)

            if not documents_url:
                logger.error(f"Could not determine documents URL for: {lot_url}")
                return []

            if progress_callback:
                progress_callback("Загрузка страницы документов.")

            logger.info(
                f"Downloading documents from: {documents_url} with params={params}"
            )

            response = self.session.get(
                documents_url, params=params, timeout=CONFIG["REQUEST_TIMEOUT"]
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "html.parser")

            # Сохраняем HTML для отладки
            with open("debug_documents_page.html", "w", encoding="utf-8") as f:
                f.write(soup.prettify())

            # Ищем документы - несколько стратегий
            document_links = []

            # Стратегия 1: Ищем по классам, которые обычно используются для документов
            document_selectors = [
                "a[href*='/file/get/']",
                "a[href*='file/get']",
                "a[href*='download']",
                "a[href*='.pdf']",
                "a[href*='.doc']",
                "a[href*='.docx']",
                "a[href*='.xls']",
                "a[href*='.xlsx']",
                "a[href*='.zip']",
                "a[href*='.rar']",
                ".document-link",
                ".file-link",
                "[class*='document'] a",
                "[class*='file'] a",
                ".cardFile",
                ".file-download",
                "tr td a",        # Ссылки в таблицах
                ".table-block a", # Ссылки в табличных блоках
            ]

            for selector in document_selectors:
                links = soup.select(selector)
                for link in links:
                    href = link.get("href")
                    if href:
                        full_url = urljoin(documents_url, href)

                        blocked, reason = self._is_blacklisted_document_url(full_url)
                        if blocked:
                            logger.debug("Skip link by blacklist (%s): %s", reason, full_url)
                            continue

                        if source_name == "tektorg.ru" and not self._is_tektorg_allowed_doc_url(full_url):
                            continue

                        # Проверяем, что это документ
                        if self._is_document_link(full_url, link):
                            name = self._get_document_name(link)
                            document_links.append({"name": name, "url": full_url})

            # Стратегия 2: Ищем все ссылки и фильтруем по расширениям
            all_links = soup.find_all("a", href=True)
            for link in all_links:
                href = link.get("href")
                if href:
                    full_url = urljoin(documents_url, href)

                    # Пропускаем, если уже есть в списке
                    if any(doc["url"] == full_url for doc in document_links):
                        continue

                    blocked, reason = self._is_blacklisted_document_url(full_url)
                    if blocked:
                        logger.debug("Skip link by blacklist (%s): %s", reason, full_url)
                        continue

                    if source_name == "tektorg.ru" and not self._is_tektorg_allowed_doc_url(full_url):
                        continue

                    if self._is_document_link(full_url, link):
                        name = self._get_document_name(link)
                        document_links.append({"name": name, "url": full_url})

            # Убираем дубликаты
            unique_links = []
            seen_urls = set()
            for doc in document_links:
                if doc["url"] not in seen_urls:
                    unique_links.append(doc)
                    seen_urls.add(doc["url"])

            logger.info(f"Found {len(unique_links)} unique document links")

            # Скачиваем документы
            downloaded_docs = []
            for i, doc in enumerate(unique_links):
                if progress_callback:
                    progress_callback(
                        f"Скачивание {i + 1}/{len(unique_links)}: {doc['name'][:30]}..."
                    )

                try:
                    doc_response = self.session.get(
                        doc["url"], timeout=CONFIG["REQUEST_TIMEOUT"], stream=True
                    )
                    doc_response.raise_for_status()

                    content = doc_response.content
                    display_name = self._normalize_filename(doc.get("name"))

                    # Пытаемся получить реальное имя файла из заголовков ответа или URL
                    real_name = self._guess_real_filename(
                        doc_response, display_name, doc["url"]
                    )

                    # Имя, которое показываем в UI
                    final_name = real_name or display_name
                    if not final_name:
                        final_name = self._normalize_filename(
                            os.path.basename(urlparse(doc["url"]).path)
                        )
                    if not final_name:
                        final_name = "document"

                    downloaded_docs.append(
                        {
                            "name": final_name,
                            "filename": real_name or final_name,
                            "content": content,
                            "size": len(content),
                            "url": doc["url"],
                            "content_type": doc_response.headers.get(
                                "content-type", ""
                            ),
                        }
                    )

                    logger.info(
                        "Successfully downloaded: display=%r, real=%r, final=%r, size=%s bytes",
                        display_name,
                        real_name,
                        final_name,
                        len(content),
                    )

                    time.sleep(0.5)  # Задержка между запросами

                except Exception as e:
                    logger.warning(f"Failed to download document {doc['name']}: {e}")
                    continue

            expanded_docs = self._expand_downloaded_documents(
                downloaded_docs,
                progress_callback=progress_callback,
            )
            if len(expanded_docs) != len(downloaded_docs):
                logger.info(
                    "Expanded documents from %s to %s after archive extraction",
                    len(downloaded_docs),
                    len(expanded_docs),
                )
            else:
                logger.info("Successfully downloaded %s documents", len(downloaded_docs))
            return expanded_docs

        except Exception as e:
            logger.error(f"Error downloading documents: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return []


    def _is_document_link(self, url, link_element):
        """Проверяет, является ли ссылка документом."""
        # Проверяем расширения файлов
        document_extensions = [
            ".pdf",
            ".doc",
            ".docx",
            ".xls",
            ".xlsx",
            ".zip",
            ".rar",
            ".txt",
            ".rtf",
            ".odt",
            ".ods",
        ]
        if any(url.lower().endswith(ext) for ext in document_extensions):
            return True

        # Проверяем ключевые слова в URL
        url_keywords = [
            "/download/",
            "/file/",
            "document",
            "file",
            "download",
            "getFile",
        ]
        if any(keyword in url.lower() for keyword in url_keywords):
            return True

        # Проверяем классы элемента
        element_classes = link_element.get("class", [])
        class_keywords = ["document", "file", "download", "cardFile", "file-download"]
        if any(
                any(keyword in str(cls).lower() for keyword in class_keywords)
                for cls in element_classes
        ):
            return True

        # Проверяем текст ссылки
        link_text = link_element.get_text(strip=True).lower()
        text_keywords = ["скачать", "документ", "файл", "doc", "pdf", "xls", "zip"]
        if any(keyword in link_text for keyword in text_keywords):
            return True

        return False

    def _get_document_name(self, link_element):
        """Извлекает имя документа из элемента ссылки."""
        # Пробуем получить имя из текста ссылки
        # tektorg: имя документа лежит в div.hAcoWe (внутри ссылки)
        try:
            name_block = link_element.select_one(".hAcoWe")
            if name_block:
                txt = name_block.get_text(" ", strip=True)
                if txt:
                    # Часто приходит "X (X)" — уберём дубль
                    if "(" in txt and txt.endswith(")"):
                        left = txt.split("(", 1)[0].strip()
                        inside = txt.rsplit("(", 1)[1].rstrip(")").strip()
                        if left and inside and left == inside:
                            txt = left
                    return txt
        except Exception:
            pass

        name = link_element.get_text(strip=True)
        if name and len(name) > 2:
            return name

        # Пробуем получить из атрибута title
        name = link_element.get("title", "")
        if name and len(name) > 2:
            return name

        # Пробуем получить из data-атрибутов
        for attr in link_element.attrs:
            if "name" in attr.lower() or "title" in attr.lower():
                name = link_element.get(attr, "")
                if name and len(name) > 2:
                    return name

        # Если ничего не нашли, генерируем имя из URL
        href = link_element.get("href", "")
        if href:
            # Извлекаем имя файла из URL
            filename = href.split("/")[-1]
            if "?" in filename:
                filename = filename.split("?")[0]
            if filename and "." in filename:
                return filename

        return f"document_{int(time.time())}"

    def search_in_documents(self, documents, keywords, progress_callback=None):
        """Ищет ключевые слова в документах."""
        results = []
        total = len(documents)

        logger.info(
            "Начат анализ документов: всего %s, ключевые слова: %s",
            total,
            ", ".join(map(str, keywords)),
        )

        for idx, doc in enumerate(documents, start=1):
            filename = self._determine_document_filename(doc)
            filename_l = filename.lower()
            doc_results = {
                "document_name": doc["name"],
                "size": doc["size"],
                "url": doc["url"],
                "matches": [],
                "match_count": 0,
            }

            if progress_callback:
                try:
                    progress_callback(
                        f"Анализ документа {idx}/{total}: {doc['name'][:60]}..."
                    )
                except Exception:
                    # Не даём GUI-ошибкам завалить анализ
                    pass

            logger.info(
                "Анализ документа %s/%s: name=%r, url=%r, filename_for_parser=%r",
                idx,
                total,
                doc.get("name"),
                doc.get("url"),
                filename,
            )

            try:
                content_text = self._extract_text_from_content(
                    doc["content"], filename, doc.get("content_type", "")
                )

                if idx == 1:
                    # для отладки: покажем фрагмент извлечённого текста первого документа
                    logger.info(
                        "Пример извлечённого текста (первые 500 символов) из %r: %r",
                        doc.get("name"),
                        content_text[:500],
                    )

                logger.debug(
                    "Извлечено %s символов текста из документа %r",
                    len(content_text),
                    doc.get("name"),
                )

                # Для Word‑документов и Excel используем более точный поиск по словам
                if (filename_l.endswith(".docx") or filename_l.endswith(".doc") or
                    filename_l.endswith((".xlsx", ".xlsm", ".xltx", ".xltm")) or
                    filename_l.endswith(".xls")):
                    matched, match_count = self._find_word_matches_in_text(
                        content_text, keywords
                    )
                    doc_type = "Word" if filename_l.endswith((".docx", ".doc")) else "Excel"
                    logger.info(
                        "Совпадения в %s‑документе %r: найдено %s, ключевые слова: %s",
                        doc_type,
                        doc.get("name"),
                        match_count,
                        matched,
                    )
                    doc_results["matches"].extend(matched)
                    doc_results["match_count"] += match_count
                else:
                    # Для остальных форматов оставляем прежний простой поиск подстроки
                    lower_text = content_text.lower()
                    for keyword in keywords:
                        keyword_lower = keyword.lower().strip()
                        if keyword_lower and keyword_lower in lower_text:
                            doc_results["matches"].append(keyword)
                            doc_results["match_count"] += 1

                    if doc_results["match_count"]:
                        logger.info(
                            "Совпадения в документе %r: найдено %s, ключевые слова: %s",
                            doc.get("name"),
                            doc_results["match_count"],
                            doc_results["matches"],
                        )

                # Добавляем контекст для первых совпадений
                if doc_results["matches"]:
                    doc_results["sample_context"] = self._get_keyword_context(
                        content_text, keywords, max_contexts=2
                    )

                results.append(doc_results)

            except Exception as e:
                logger.warning(f"Error analyzing document {doc['name']}: {e}")
                results.append(doc_results)
                continue

        logger.info(
            "Анализ документов завершён. Документов: %s, с совпадениями: %s",
            len(results),
            sum(1 for r in results if r["matches"]),
        )

        # Сортируем по количеству совпадений
        results.sort(key=lambda x: x["match_count"], reverse=True)
        return results

    def _determine_document_filename(self, doc: dict) -> str:
        """
        Пытается определить реальное имя файла:
        - сначала используем то, что вывели в UI (doc['name'])
        - если расширения нет, пробуем взять basename из URL
        """
        name = (doc.get("filename")
                or doc.get("name")
                or "").strip()
        if "." in os.path.basename(name):
            return name

        url = doc.get("url") or ""
        try:
            path = urlparse(url).path
            candidate = os.path.basename(path)
            if candidate:
                return candidate
        except Exception:
            pass

        # Фолбэк — оставляем оригинальное имя или генерируем
        return name or f"document_{int(time.time())}"

    def _find_word_matches_in_text(self, text: str, keywords, min_len: int = 2):
        """
        Более аккуратный поиск по словам в тексте, в первую очередь для Word‑документов.

        - Использует регулярные выражения с границами слов (\b) и учётом юникода.
        - Игнорирует регистр.
        - Для каждого ключевого слова увеличивает счётчик на количество найденных вхождений.
        """
        if not text:
            return [], 0

        text_normalized = " ".join(str(text).split())
        total_count = 0
        matched_keywords = []

        for raw_kw in keywords:
            kw = (raw_kw or "").strip()
            if not kw or len(kw) < min_len:
                continue

            # Экранируем спецсимволы, чтобы искать именно текст, а не паттерн regex
            pattern = r"\b" + re.escape(kw) + r"\b"
            matches = re.findall(pattern, text_normalized, flags=re.IGNORECASE)
            if matches:
                matched_keywords.append(raw_kw)
                total_count += len(matches)

        return matched_keywords, total_count

    def _extract_text_from_content(self, content, filename, content_type):
        """Извлекает текст из содержимого документа с учётом формата."""
        filename_l = (filename or "").lower()
        ctype = (content_type or "").lower()

        try:
            # 0. Если это ZIP‑контейнер (OOXML: docx/xlsx и т.п.) — определяем тип по содержимому
            #    Независимо от расширения файла, так как на сайте часто путают форматы.
            if content[:2] == b"PK":
                # Проверяем, что это за OOXML формат
                try:
                    with zipfile.ZipFile(BytesIO(content)) as zf:
                        file_list = zf.namelist()
                        # XLSX содержит xl/workbook.xml или xl/sharedStrings.xml
                        if any(f.startswith("xl/") for f in file_list):
                            text = self._extract_text_from_xlsx(content)
                            if text:
                                return text
                        # DOCX содержит word/document.xml
                        elif any(f.startswith("word/") for f in file_list):
                            text = self._extract_text_from_docx(content)
                            if text:
                                return text
                except Exception:
                    # Если не удалось определить, пробуем как DOCX (более частый случай)
                    text = self._extract_text_from_docx(content)
                    if text:
                        return text

            # 1. Обычный текст
            if filename_l.endswith(".txt") or "text/" in ctype:
                return content.decode("utf-8", errors="ignore")

            # 2. DOCX
            if (
                    filename_l.endswith(".docx")
                    or "officedocument.wordprocessingml" in ctype
            ):
                text = self._extract_text_from_docx(content)
                if text:
                    return text

            # 2.1. Старые DOC (часто RTF или HTML внутри .doc)
            if (
                    filename_l.endswith(".doc")
                    and not filename_l.endswith(".docx")
            ) or "msword" in ctype:
                text = self._extract_text_from_doc(content)
                if text:
                    return text

            # 3. XLSX / XLSM
            if filename_l.endswith((".xlsx", ".xlsm", ".xltx", ".xltm")) or "spreadsheetml" in ctype:
                text = self._extract_text_from_xlsx(content)
                if text:
                    return text

            # 4. Старые XLS
            if filename_l.endswith(".xls") or "ms-excel" in ctype:
                text = self._extract_text_from_xls(content)
                if text:
                    return text

            # Если это PDF, а нормального PDF-парсера нет — не пытаемся decode(), иначе будет мусор
            if content[:4] == b"%PDF":
                return ""

            # Если много нулевых байт — это бинарь, не текст
            if b"\x00" in content[:4096]:
                return ""

            # 5. Всё остальное (PDF, DOC, RTF и т.п.) — грубый универсальный метод
            # Сначала пробуем UTF-8, чтобы не сломать кириллицу, даже если формат распознан неверно.
            text = content.decode("utf-8", errors="ignore")
            text = re.sub(r"[^\x20-\x7E\xC0-\xFF\n\r\t]", " ", text)
            text = re.sub(r"\s+", " ", text)
            return text

        except Exception as e:
            logger.warning(f"Could not extract text from {filename}: {e}")
            return ""


    def _extract_text_from_docx(self, content: bytes) -> str:
        """Извлекает текст из .docx с помощью python-docx."""
        # 1) Основной путь — через python-docx (если установлен)
        if Document is not None:
            try:
                doc = Document(BytesIO(content))
                parts = []

                # Параграфы
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        parts.append(text)

                # Таблицы
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " ".join(
                            (cell.text or "").strip() for cell in row.cells if cell.text
                        )
                        if row_text:
                            parts.append(row_text)

                return "\n".join(parts)
            except Exception as e:
                logger.warning(f"Не удалось извлечь текст из DOCX через python-docx: {e}")

        # 2) Запасной путь — разбираем .docx как ZIP и достаём word/document.xml
        try:
            with zipfile.ZipFile(BytesIO(content)) as zf:
                with zf.open("word/document.xml") as doc_xml:
                    xml_bytes = doc_xml.read()
            xml_text = xml_bytes.decode("utf-8", errors="ignore")
            # Убираем XML/HTML-теги и декодируем сущности (&amp; и т.п.)
            xml_text = html.unescape(re.sub(r"<[^>]+>", " ", xml_text))
            xml_text = re.sub(r"\s+", " ", xml_text)
            return xml_text.strip()
        except Exception as e:
            logger.warning(f"Не удалось извлечь текст из DOCX через zipfile: {e}")
            return ""

    def _guess_real_filename(self, response, display_name: str | None, url: str) -> str:
        """
        Определяет реальное имя файла по следующему приоритету:
        1) filename / filename* в заголовке Content-Disposition
        2) basename из URL (если есть расширение)
        3) display_name (как в ссылке на сайте)
        """
        # 1. Content-Disposition
        cd = response.headers.get("content-disposition", "") or ""
        cd_lower = cd.lower()
        filename = ""

        # filename* (RFC 5987, например: filename*=UTF-8''%D0%90%D0%BA%D1%82.docx)
        if "filename*=" in cd_lower:
            try:
                part = cd.split("filename*=", 1)[1].split(";", 1)[0].strip()
                if part.lower().startswith("utf-8''"):
                    part = part[8:]
                part = part.strip('";')
                filename = unquote(part)
            except Exception:
                filename = ""

        # обычный filename="name.ext"
        if not filename and "filename=" in cd_lower:
            try:
                part = cd.split("filename=", 1)[1].split(";", 1)[0].strip()
                filename = part.strip('";')
            except Exception:
                filename = ""

        # 2. basename из URL, если в нём есть расширение
        if not filename:
            try:
                path = urlparse(url).path
                base = os.path.basename(path)
                if "." in base:
                    filename = base
            except Exception:
                filename = ""

        # 3. display_name как есть
        if not filename:
            filename = (display_name or "").strip()

        return self._normalize_filename(filename or "document")

    def _normalize_filename(self, name: str | None) -> str:
        """Приводит имя файла к удобному виду и исправляет типичный mojibake."""
        if not name:
            return ""

        cleaned = (name or "").strip().strip('"').replace("\\", "/")
        cleaned = cleaned.split("/")[-1]
        if not cleaned:
            return ""

        if self._looks_like_mojibake(cleaned):
            for encoding in ("utf-8", "cp1251"):
                try:
                    cleaned = cleaned.encode("latin-1").decode(encoding)
                    break
                except Exception:
                    continue

        return cleaned.strip()

    def _looks_like_mojibake(self, text: str) -> bool:
        """Пытается определить, искажено ли имя файла (Ð, Ñ, Ã и т.п.)."""
        if not text:
            return False
        mojibake_markers = ("Ð", "Ñ", "Ã", "Ò", "Â")
        return any(marker in text for marker in mojibake_markers)

    def _extract_text_from_doc(self, content: bytes) -> str:
        """
        Пытается извлечь текст из старого .doc.

        Частые варианты на гос-сайтах:
        - RTF, сохранённый как .doc
        - HTML, сохранённый как .doc
        - бинарный DOC, где текстовые фрагменты всё равно можно частично вытащить декодированием.
        """
        header = content[:2048]

        # 0) Если установлен Microsoft Word и доступен win32com — сначала пробуем извлечь текст через него.
        #    Это самый надёжный способ для «настоящих» бинарных DOC, независимо от сигнатуры.
        if self._can_use_win32_word():
            text = self._extract_text_from_doc_with_word(content)
            if text:
                return text

        # 1) RTF внутри .doc
        if b"{\\rtf" in header.lower():
            try:
                # В RTF для русских документов часто используется cp1251
                txt = content.decode("cp1251", errors="ignore")

                # Преобразуем последовательности \'hh в символы cp1251
                def _rtf_hex_to_char(match):
                    try:
                        byte_val = int(match.group(1), 16)
                        return bytes([byte_val]).decode("cp1251", errors="ignore")
                    except Exception:
                        return ""

                txt = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: _rtf_hex_to_char(m), txt)

                # Убираем управляющие последовательности RTF и фигурные скобки
                txt = re.sub(r"\\[a-zA-Z]+\d*", " ", txt)  # команды \b, \par, \fs20 и т.п.
                txt = re.sub(r"[{}]", " ", txt)
                txt = html.unescape(txt)
                txt = re.sub(r"\s+", " ", txt)
                return txt.strip()
            except Exception as e:
                logger.warning(f"Не удалось извлечь текст из RTF внутри DOC: {e}")

        # 2) HTML внутри .doc
        if b"<html" in header.lower() or b"<body" in header.lower():
            try:
                # Пробуем cp1251, затем UTF-8
                try:
                    html_text = content.decode("cp1251")
                except UnicodeDecodeError:
                    html_text = content.decode("utf-8", errors="ignore")

                soup = BeautifulSoup(html_text, "html.parser")
                txt = soup.get_text(separator=" ", strip=True)
                txt = re.sub(r"\s+", " ", txt)
                return txt.strip()
            except Exception as e:
                logger.warning(f"Не удалось извлечь HTML-текст из DOC: {e}")

        # 3) Бинарный DOC — грубая попытка вытащить видимый текст
        try:
            # Сначала cp1251, чтобы не сломать кириллицу
            txt = content.decode("cp1251", errors="ignore")
            # Оставляем только печатные символы и пробелы
            txt = re.sub(r"[^\x20-\x7E\xC0-\xFF\n\r\t]", " ", txt)
            txt = re.sub(r"\s+", " ", txt)
            return txt.strip()
        except Exception as e:
            logger.warning(f"Не удалось извлечь текст из бинарного DOC: {e}")
            return ""

    def _can_use_win32_word(self) -> bool:
        return (
                sys.platform.startswith("win")
                and win32_client is not None
                and win32_constants is not None
                and pythoncom is not None
        )

    def _extract_text_from_doc_with_word(self, content: bytes) -> str:
        """Использует установленный Microsoft Word (через win32com) для извлечения текста из .doc."""
        temp_doc = None
        temp_txt = None
        word = None
        try:
            pythoncom.CoInitialize()
            temp = tempfile.NamedTemporaryFile(delete=False, suffix=".doc")
            temp.write(content)
            temp_doc = temp.name
            temp.close()

            temp_txt = temp_doc + ".txt"

            word = win32_client.Dispatch("Word.Application")  # type: ignore
            word.Visible = False

            doc = word.Documents.Open(temp_doc)  # type: ignore
            wd_format_text = getattr(win32_constants, "wdFormatText", 2)
            doc.SaveAs(temp_txt, FileFormat=wd_format_text)  # type: ignore
            doc.Close(False)  # type: ignore

            with open(temp_txt, "rb") as f:
                raw = f.read()

            encodings_to_try = [
                "utf-8",
                locale.getpreferredencoding(False) or "cp1251",
                "cp1251",
            ]
            text = ""
            for enc in encodings_to_try:
                try:
                    text = raw.decode(enc)
                    break
                except Exception:
                    continue
            if not text:
                text = raw.decode("latin-1", errors="ignore")

            return " ".join(text.split())

        except Exception as e:
            logger.warning(f"Не удалось извлечь текст из DOC через Word COM: {e}")
            return ""
        finally:
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            if pythoncom is not None:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass
            for path in (temp_doc, temp_txt):
                if not path:
                    continue
                try:
                    os.remove(path)
                except Exception:
                    pass

    def _extract_text_from_xlsx(self, content: bytes) -> str:
        """Извлекает текст из .xlsx/.xlsm с помощью openpyxl."""
        if load_workbook is None:
            logger.warning("openpyxl не установлен, невозможно извлечь текст из XLSX")
            return ""

        try:
            wb = load_workbook(BytesIO(content), data_only=True, read_only=True)
            parts = []
            total_cells = 0

            for ws in wb.worksheets:
                sheet_name = ws.title
                logger.debug(f"Обработка листа Excel: {sheet_name}")
                
                for row in ws.iter_rows():
                    row_vals = []
                    for cell in row:
                        v = cell.value
                        if v is None:
                            continue
                        # Преобразуем значение в строку, обрабатывая даты и числа
                        if isinstance(v, (int, float)):
                            row_vals.append(str(v))
                        elif isinstance(v, str):
                            row_vals.append(v.strip())
                        else:
                            row_vals.append(str(v))
                        total_cells += 1
                    
                    if row_vals:
                        parts.append(" ".join(row_vals))

            result = "\n".join(parts)
            logger.debug(f"Извлечено {len(parts)} строк из {len(wb.worksheets)} листов Excel, всего ячеек: {total_cells}")
            return result
        except Exception as e:
            logger.warning(f"Не удалось извлечь текст из XLSX: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return ""

    def _extract_text_from_xls(self, content: bytes) -> str:
        """Извлекает текст из .xls с помощью xlrd (если установлен)."""
        if xlrd is None:
            logger.warning("xlrd не установлен, невозможно извлечь текст из XLS")
            return ""
        try:
            book = xlrd.open_workbook(file_contents=content)
            parts = []
            total_cells = 0

            for sheet in book.sheets():
                sheet_name = sheet.name
                logger.debug(f"Обработка листа Excel (XLS): {sheet_name}")
                
                for rx in range(sheet.nrows):
                    row_vals = sheet.row_values(rx)
                    # Преобразуем значения в строки, пропуская пустые
                    row_vals = [str(v).strip() for v in row_vals if v not in ("", None) and str(v).strip()]
                    if row_vals:
                        parts.append(" ".join(row_vals))
                        total_cells += len(row_vals)
            
            result = "\n".join(parts)
            logger.debug(f"Извлечено {len(parts)} строк из {len(book.sheets())} листов Excel (XLS), всего ячеек: {total_cells}")
            return result
        except Exception as e:
            logger.warning(f"Не удалось извлечь текст из XLS: {e}")
            import traceback
            logger.debug(traceback.format_exc())
            return ""



    def _get_keyword_context(self, text, keywords, max_contexts=2, context_length=100):
        """Извлекает контекст вокруг найденных ключевых слов."""
        contexts = []
        text_lower = text.lower()

        for keyword in keywords:
            keyword_lower = keyword.lower().strip()
            if not keyword_lower:
                continue

            start = 0
            while len(contexts) < max_contexts:
                pos = text_lower.find(keyword_lower, start)
                if pos == -1:
                    break

                start_context = max(0, pos - context_length)
                end_context = min(len(text), pos + len(keyword_lower) + context_length)

                context = text[start_context:end_context]
                context = " ".join(context.split())
                contexts.append(f"...{context}...")

                start = pos + len(keyword_lower)

        return contexts[:max_contexts]

